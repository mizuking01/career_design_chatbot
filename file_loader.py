from langchain_core.documents import Document 
from langchain.document_loaders import PyMuPDFLoader, TextLoader  
from docx import Document as DocxDocument

def load_docx(docx_path):
    docx = DocxDocument(docx_path)
    content = "\n".join([para.text for para in docx.paragraphs])
    return [Document(page_content=content, metadata={"source": docx_path})]  # ← 修正！

def load_pdf(pdf_path):
    loader = PyMuPDFLoader(pdf_path)
    return loader.load()  # PyMuPDFLoader は Document を返すのでそのままでOK！

def load_text(text_path):
    loader = TextLoader(text_path)
    return loader.load()  # TextLoader も Document を返すのでOK！
